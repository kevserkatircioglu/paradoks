import io
import re
import zipfile
import docx
from pathlib import PurePosixPath

import requests
from bs4 import BeautifulSoup


HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 "
        "(compatible; standards-crawler/1.0)"
    )
}

TIMEOUT = 30


def fetch_and_read(
    url: str,
    requested_code: str | None = None
) -> str | None:
    """
    URL'deki dokümanı indirir ve metin içeriğini döndürür.
    requested_code verilirse ZIP içinden doğru belgenin 
    filtrelenmesini sağlar.
    """

    try:
        resp = requests.get(
            url,
            headers=HEADERS,
            timeout=TIMEOUT,
        )

    except requests.RequestException:
        return None

    if resp.status_code != 200:
        return None

    content_type = (
        resp.headers.get(
            "Content-Type",
            "",
        ).lower()
    )

    lower_url = (
        url.lower()
    )

    # -----------------------------------------------------
    # PDF
    # -----------------------------------------------------
    if (
        lower_url.endswith(".pdf")
        or "application/pdf"
        in content_type
    ):
        return _read_pdf(
            resp.content
        )

    # -----------------------------------------------------
    # DOCX
    # -----------------------------------------------------
    if (
        lower_url.endswith(".docx")
        or "wordprocessingml"
        in content_type
    ):
        return _read_docx(
            resp.content
        )

    # -----------------------------------------------------
    # ZIP
    # -----------------------------------------------------
    if (
        lower_url.endswith(".zip")
        or "application/zip"
        in content_type
        or "application/x-zip-compressed"
        in content_type
    ):
        return _read_zip(
            resp.content,
            requested_code
        )

    # -----------------------------------------------------
    # HTML
    # -----------------------------------------------------
    return _read_html(
        resp.text
    )


def _read_pdf(
    raw_bytes: bytes,
) -> str:
    """
    PDF içeriğini sayfa sayfa okuyarak metne çevirir.
    """

    import pdfplumber

    text_parts: list[str] = []

    with pdfplumber.open(
        io.BytesIO(
            raw_bytes
        )
    ) as pdf:

        for page in pdf.pages:
            text_parts.append(
                page.extract_text()
                or ""
            )

    return "\n".join(
        text_parts
    )


def _read_docx(
    raw_bytes: bytes,
) -> str:
    """
    DOCX dosyasındaki paragraph metinlerini ve tabloları çıkarır.
    Bozuk dosyalarda kodun çökmesini engeller.
    """

    try:
        doc = docx.Document(
            io.BytesIO(
                raw_bytes
            )
        )
        
        full_text = []

        # 1. Paragrafları oku
        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        # 2. Tablo içindeki hücreleri oku (Standardın asıl içeriği genelde buradadır)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())

        return "\n".join(full_text)

    except Exception as e:
        print(f"[FETCHER HATA] DOCX okunamadı veya bozuk: {e}")
        return ""


def _docx_sort_key(
    filename: str,
) -> tuple[int, str]:
    """
    3GPP'nin çok parçalı DOCX paketlerini doğru sıraya dizer.
    """

    basename = (
        filename
        .replace("\\", "/")
        .split("/")[-1]
    )

    match = re.search(
        r"_(\d+)_",
        basename,
    )

    if match:
        return (
            int(
                match.group(1)
            ),
            basename.lower(),
        )

    return (
        999999,
        basename.lower(),
    )


def _read_zip(
    raw_bytes: bytes,
    requested_code: str | None = None
) -> str:
    """
    ZIP içindeki TÜM geçerli DOCX dosyalarını okur.
    macOS metadata dosyalarını eler ve hedef koda uymayan alakasız dosyaları reddeder.
    """

    with zipfile.ZipFile(
        io.BytesIO(
            raw_bytes
        )
    ) as archive:

        docx_names = [
            name
            for name
            in archive.namelist()
            if name.lower().endswith(".docx")
            and not name.startswith("__MACOSX/")
            and not name.split("/")[-1].startswith("~$")
            and not name.split("/")[-1].startswith("._")  # EKLENDİ: macOS metadata gizli dosyaları elendi
        ]

        if not docx_names:
            return ""

        # EKLENDİ: Eğer istenen bir kod varsa (örn: TS 23.366),
        # sadece isminde '23366' geçen dosyaları filtrele (Yanlış belge okumayı engeller)
        if requested_code:
            norm_code = re.sub(r"\D", "", requested_code)
            if norm_code:
                matched_names = [n for n in docx_names if norm_code in n.split("/")[-1]]
                if matched_names:
                    docx_names = matched_names

        docx_names.sort(
            key=_docx_sort_key
        )

        text_parts: list[str] = []

        for docx_name in docx_names:
            try:
                with archive.open(
                    docx_name
                ) as file:
                    raw_docx = (
                        file.read()
                    )

                extracted_text = (
                    _read_docx(
                        raw_docx
                    )
                ).strip()

                if not extracted_text:
                    continue

                text_parts.append(
                    extracted_text
                )

            except Exception as error:
                print(
                    "[FETCHER] ZIP içindeki DOCX okunamadı:",
                    docx_name,
                    "|",
                    error,
                )
                continue

        return "\n\n".join(
            text_parts
        )


def _read_html(
    html: str,
) -> str:
    """
    HTML sayfasındaki okunabilir metni çıkarır.
    """

    soup = BeautifulSoup(
        html,
        "html.parser",
    )

    return soup.get_text(
        separator="\n",
        strip=True,
    )